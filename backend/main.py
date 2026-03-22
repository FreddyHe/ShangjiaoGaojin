from __future__ import annotations
import os
import json
from pathlib import Path
from typing import List, Optional, Dict, Set, Tuple
import re
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from pypdf import PdfReader
from dotenv import load_dotenv
import markdown
# from xhtml2pdf import pisa
from io import BytesIO
# from reportlab.pdfbase import pdfmetrics
# from reportlab.pdfbase.ttfonts import TTFont
import asyncio
import aiohttp
from urllib.parse import quote
from pypinyin import lazy_pinyin
import requests
from bs4 import BeautifulSoup
import time
import shutil

COMMON_SURNAMES = set("赵钱孙李周吴郑王冯陈褚卫蒋沈韩杨朱秦尤许何吕施张孔曹严华金魏陶姜戚谢邹喻柏水窦章云苏潘葛奚范彭郎鲁韦昌马苗凤花方俞任袁柳酆鲍史唐费廉岑薛雷贺倪汤滕殷罗毕郝安常乐于时傅皮卞齐康伍余元卜顾孟平黄和穆萧尹姚邵湛汪祁毛禹狄米贝明臧计伏成戴谈宋茅庞熊纪舒屈项祝董梁杜阮蓝闵席季麻强贾路娄危江童颜郭梅盛林刁钟徐邱骆高夏蔡田樊胡凌霍虞万支柯昝管卢莫经房裘缪干解应宗丁宣邓郁单杭洪包诸左石崔吉钮龚程嵇邢滑裴陆荣翁荀羊於惠甄曲家封芮羿储晋汲邴糜松井段富巫乌焦巴弓牧隗山谷车侯宓蓬全郗班仰秋仲伊宫宁仇栾暴甘钭厉戎祖武符刘景詹束龙叶幸司韶郜黎蓟薄印宿白怀蒲台从鄂索咸籍赖卓蔺屠蒙池乔阴郁胥能苍双闻莘党翟谭贡劳逄姬申扶堵冉宰郦雍却璩桑桂濮牛寿通边扈燕冀郟浦尚农温别庄晏柴瞿阎充慕连茹习宦艾鱼容向古易慎戈廖庾终暨居衡步都耿满弘匡国文寇广禄阙东欧沃利蔚越夔隆师巩厍聂晁勾敖融冷訾辛阚那简饶空曾毋沙乜养鞠须丰巢关蒯相查后荆红游竺权逯盖益桓公")

load_dotenv()

# --- 核心算法工具函数 ---

def levenshtein_ratio(s1: list, s2: list) -> float:
    """
    计算两个序列(字符串或列表)的相似度比率 (0~1)。
    使用标准 Levenshtein 距离算法。
    """
    if not s1 or not s2:
        return 0.0
    
    len1, len2 = len(s1), len(s2)
    max_len = max(len1, len2)
    if max_len == 0:
        return 1.0

    # 初始化矩阵
    matrix = [[0] * (len2 + 1) for _ in range(len1 + 1)]
    for i in range(len1 + 1):
        matrix[i][0] = i
    for j in range(len2 + 1):
        matrix[0][j] = j

    # 动态规划计算距离
    for i in range(1, len1 + 1):
        for j in range(1, len2 + 1):
            cost = 0 if s1[i - 1] == s2[j - 1] else 1
            matrix[i][j] = min(
                matrix[i - 1][j] + 1,      # 删除
                matrix[i][j - 1] + 1,      # 插入
                matrix[i - 1][j - 1] + cost # 替换
            )
    
    distance = matrix[len1][len2]
    return 1.0 - (distance / max_len)

def calculate_mixed_similarity(candidate: str, db_name: str) -> float:
    """
    核心评分逻辑：(字形相似度 + 拼音相似度) / 2
    """
    # 1. 计算文字相似度 (y)
    text_similarity_y = levenshtein_ratio(list(candidate), list(db_name))
    
    # 如果文字完全一样，直接返回 1.0 (节省计算)
    if text_similarity_y == 1.0:
        return 1.0

    # 2. 计算拼音相似度 (x)
    # lazy_pinyin 会把 "程仕军" 变成 ['cheng', 'shi', 'jun']
    pinyin_cand = lazy_pinyin(candidate)
    pinyin_db = lazy_pinyin(db_name)
    
    # 优化：拼接为字符串，再转为字符列表做字符级比较
    # 这样 "chengshijun" 和 "chengsijun" 的编辑距离只有1（删除 h），而不是列表级的1
    pinyin_str_cand = ''.join(pinyin_cand)
    pinyin_str_db = ''.join(pinyin_db)
    pinyin_similarity_x = levenshtein_ratio(list(pinyin_str_cand), list(pinyin_str_db))

    # 3. 按照公式取平均值
    final_score = (pinyin_similarity_x + text_similarity_y) / 2
    
    return final_score

def extract_candidates(text: str) -> Set[str]:
    """
    步骤1：利用【百家姓】从素材中提取潜在人名
    """
    candidates = set()
    length = len(text)
    
    for i in range(length):
        # 【逻辑1：百家姓过滤】如果不以常见姓氏开头，直接跳过
        if text[i] not in COMMON_SURNAMES:
            continue
            
        # 提取长度为 2 和 3 的片段 (大多数中文名长度)
        # 如果需要支持复姓或4字名，可以加一个 4
        for name_len in [2, 3]: 
            if i + name_len <= length:
                segment = text[i : i + name_len]
                # 再次确认片段里是否全是中文（避免 "李a" 这种情况）
                if re.match(r'^[\u4e00-\u9fa5]+$', segment):
                    candidates.add(segment)
    
    return candidates

def extract_context_around_name(text: str, name: str, context_length: int = 100) -> str:
    """Extract context around a name in the text"""
    index = text.find(name)
    if index == -1:
        return ""
    
    start = max(0, index - context_length)
    end = min(len(text), index + len(name) + context_length)
    return text[start:end]
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
OUTLINES_DIR = DATA_DIR / "outlines"
ARTICLES_DIR = DATA_DIR / "articles"
TYPES_REGISTRY_FILE = DATA_DIR / "types_registry.json"
SETTINGS_FILE = DATA_DIR / "settings.json"


def build_master_system_prompt(outline_sections: List[OutlineSection], type_name: str, materials: str, knowledge_base: Optional[Dict[str, str]] = None) -> str:
    """构建包含大纲、素材、知识库的完整系统提示词"""
    
    # 构造 payload，保持与 generate 接口一致的结构
    payload = {
        "type": type_name,
        "outline": [_section_to_dict(s) for s in outline_sections],
        "materials": materials,
    }
    
    # 基础指令
    base_instruction = (
        "根据类型与统一大纲以及素材，撰写或修改一篇完整的中文新闻稿用于发布在公众号平台，语气专业、结构严谨、可直接发布。"
        "输出不要包含任何额外的文字，只输出新闻稿内容。\n"
    )
    
    # 知识库指令
    kb_instruction = ""
    if knowledge_base:
        payload["knowledge_base"] = knowledge_base
        kb_instruction = (
            "我提供了一个'knowledge_base'（人物知识库）。如果在撰写过程中提到了知识库中的人物（如演讲嘉宾、获奖者等），"
            "必须在第一次提及时，自然地融入其头衔、学历、职位和研究兴趣等信息。\n"
            "例如：'XXX教授（国家重点研发计划项目首席科学家、同济大学教授）在会上表示...'\n"
            "不要生硬地罗列简历，要结合语境自然插入。\n"
        )
        
    # 格式指令
    format_instruction = (
        "输出的新闻稿应该支持在建议放置图片或者视频素材的部分用特殊形式表示，比如【此处建议添加xxx图片素材】。\n"
        "遵循建议的层级大纲（包含children层级）和描述性要点、字数建议，保持专业和现代化表达，以实现宣传目的。"
    )
    
    # 组合最终 Prompt
    full_prompt = (
        f"{base_instruction}"
        f"{kb_instruction}"
        f"{format_instruction}"
        f"输入：\n{json.dumps(payload, ensure_ascii=False)}"
    )
    
    return full_prompt


app = FastAPI()

class Settings(BaseModel):
    system_name: str = "新闻稿智能体"
    logo_url: Optional[str] = None
    copyright_text: str = "© 2024 Press Release Assistant"
    openai_api_key: Optional[str] = None
    openai_base_url: Optional[str] = None
    openai_model: str = "gpt-4o-mini"

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class TypeItem(BaseModel):
    id: str
    name: str
    description: Optional[str] = None

class OutlineSection(BaseModel):
    title: str
    bullets: Optional[List[str]] = None
    children: Optional[List[OutlineSection]] = None
    description: Optional[str] = None
    word_count: Optional[str] = None
    granularity: Optional[str] = None

class OutlineMeta(BaseModel):
    is_favorite: bool = False
    created_at: float = 0
    title: Optional[str] = None

class Outline(BaseModel):
    type_id: str
    type_name: str
    template_id: Optional[str] = "default"
    sections: List[OutlineSection]
    meta: Optional[OutlineMeta] = None

class Article(BaseModel):
    id: str
    title: str
    type_id: str
    type_name: str
    template_id: str
    content: str
    materials: str
    created_at: float
    parent_id: Optional[str] = None
    trigger_query: Optional[str] = None

class GenerateRequest(BaseModel):
    type_id: str
    outline: Outline
    materials: str
    title: Optional[str] = None
    knowledge_base: Optional[Dict[str, str]] = None

class ModifyRequest(BaseModel):
    article_id: str
    user_query: str
    knowledge_base: Optional[Dict[str, str]] = None


class TypeWithStats(TypeItem):
    template_count: int

class ClassifyRequest(BaseModel):
    texts: List[str]

class PersonMatchRequest(BaseModel):
    materials: str
    people_db: Dict[str, str]

class SimilarName(BaseModel):
    name: str
    similarity: float

class PersonMatchResult(BaseModel):
    db_name: str
    material_name: str
    db_info: str
    match_type: str  # 'exact' or 'fuzzy'
    similarity: float
    similar_names: Optional[List[SimilarName]] = None  # List of similar names with similarity scores

class ConflictCheckRequest(BaseModel):
    name: str
    db_info: str
    material_context: str  # Extracted context around the name

class ConflictCheckResult(BaseModel):
    has_conflict: bool
    material_info: Optional[str] = None
    reason: Optional[str] = None
    error: Optional[str] = None
    confidence: Optional[float] = None

GENERIC_OUTLINE_SECTIONS = [
    {"title": "标题", "bullets": ["简洁明了，突出核心事件"]},
    {"title": "导语", "bullets": ["概括新闻要素（谁、什么、何时、何地、为何）"]},
    {"title": "主体内容", "bullets": ["详细描述事件背景、过程及重要细节"]},
    {"title": "结语", "bullets": ["总结意义，展望未来"]},
    {"title": "关于我们", "bullets": ["公司或机构简介"]}
]

def ensure_dirs():
    OUTLINES_DIR.mkdir(parents=True, exist_ok=True)
    ARTICLES_DIR.mkdir(parents=True, exist_ok=True)

def load_settings() -> Settings:
    if not SETTINGS_FILE.exists():
        # Default from env if file missing
        return Settings(
            openai_api_key=os.getenv("OPENAI_API_KEY"),
            openai_base_url=os.getenv("OPENAI_BASE_URL"),
            openai_model=os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        )
    try:
        data = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        return Settings(**data)
    except Exception:
        return Settings()

def save_settings(s: Settings):
    ensure_dirs()
    SETTINGS_FILE.write_text(json.dumps(s.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")

def outline_path(type_id: str, template_id: Optional[str] = None) -> Path:
    tid = template_id or "default"
    if tid == "default":
        return OUTLINES_DIR / f"{type_id}.json"
    return OUTLINES_DIR / f"{type_id}__{tid}.json"

def load_types() -> List[TypeItem]:
    if not TYPES_REGISTRY_FILE.exists():
        raise HTTPException(status_code=500, detail="types_registry.json missing")
    data = json.loads(TYPES_REGISTRY_FILE.read_text(encoding="utf-8"))
    return [TypeItem(**t) for t in data.get("types", [])]

def save_type(t: TypeItem):
    data = {"types": [x.model_dump() for x in load_types()]}
    if any(x["id"] == t.id for x in data["types"]):
        raise HTTPException(status_code=400, detail="type id exists")
    data["types"].append(t.model_dump())
    TYPES_REGISTRY_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def get_client() -> OpenAI:
    settings = load_settings()
    key = settings.openai_api_key or os.getenv("OPENAI_API_KEY")
    base_url = settings.openai_base_url or os.getenv("OPENAI_BASE_URL")
    
    if not key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY not set")
    if base_url:
        return OpenAI(api_key=key, base_url=base_url)
    return OpenAI(api_key=key)

def get_article_history(article_id: str) -> List[Article]:
    """获取文章的完整修改历史，包括原文章及其所有子文章"""
    history: List[Article] = []
    visited: Set[str] = set()  # 用于防止循环引用
    
    def add_article(article_id: str):
        """递归添加文章及其所有子文章"""
        if article_id in visited:
            return
        
        article_path = ARTICLES_DIR / f"{article_id}.json"
        if not article_path.exists():
            return
        
        try:
            article_data = json.loads(article_path.read_text(encoding="utf-8"))
            article = Article(**article_data)
            history.append(article)
            visited.add(article_id)
            
            # 递归获取所有子文章
            for p in ARTICLES_DIR.glob("*.json"):
                try:
                    child_data = json.loads(p.read_text(encoding="utf-8"))
                    if child_data.get("parent_id") == article_id:
                        add_article(child_data["id"])
                except Exception as e:
                    print(f"Error loading child article {p.name}: {e}")
        except Exception as e:
            print(f"Error loading article {article_id}: {e}")
    
    # 从给定的article_id开始递归添加
    add_article(article_id)
    
    # 按创建时间排序（最早的版本在前）
    history.sort(key=lambda x: x.created_at)
    return history

def read_txt_bytes(b: bytes) -> str:
    try:
        return b.decode("utf-8")
    except UnicodeDecodeError:
        return b.decode("gb18030", errors="ignore")

def read_docx_bytes(b: bytes) -> str:
    import uuid
    tmp = BASE_DIR / f"_tmp_upload_{uuid.uuid4()}.docx"
    tmp.write_bytes(b)
    doc = Document(str(tmp))
    text = "\n".join([p.text for p in doc.paragraphs])
    try:
        tmp.unlink()
    except Exception:
        pass
    return text

def read_pdf_bytes(b: bytes) -> str:
    reader = PdfReader(BytesIO(b))
    pages = []
    for p in reader.pages:
        t = p.extract_text() or ""
        pages.append(t)
    return "\n".join(pages)

def _safe_json_object(content: str) -> Optional[dict]:
    try:
        return json.loads(content)
    except Exception:
        pass
    try:
        start = content.find('{')
        end = content.rfind('}')
        if start != -1 and end != -1 and end > start:
            return json.loads(content[start:end+1])
    except Exception:
        return None

def _safe_json_array(content: str) -> Optional[List[dict]]:
    try:
        data = json.loads(content)
        if isinstance(data, list):
            return data
        if isinstance(data, dict):
            for k in ("sections", "items", "data"):
                v = data.get(k)
                if isinstance(v, list):
                    return v
        return None
    except Exception:
        pass
    try:
        start = content.find('[')
        end = content.rfind(']')
        if start != -1 and end != -1 and end > start:
            return json.loads(content[start:end+1])
    except Exception:
        return None

def classify_type_for_texts(texts: List[str], types: List[TypeItem]) -> TypeItem:
    client = get_client()
    settings = load_settings()
    model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    choices = [f"{t.id}:{t.name}" for t in types]
    joined = "\n\n".join(texts)[:8000]
    prompt = (
        "你是新闻稿分类助手。根据以下新闻稿合集内容，选择最匹配的已有类型，"
        "仅返回一个JSON对象，字段为id和name，id必须来自候选列表。候选："
        + ", ".join(choices)
        + "。内容：\n" + joined
    )
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            response_format={"type": "json_object"}
        )
        content = resp.choices[0].message.content or "{}"
        data = _safe_json_object(content) or {"id": types[0].id, "name": types[0].name}
        return TypeItem(**data)
    except Exception as e:
        print(e)
        return types[0]

def _sections_from_json(arr: List[dict]) -> List[OutlineSection]:
    out: List[OutlineSection] = []
    for x in arr or []:
        children = _sections_from_json(x.get("children", []) if isinstance(x.get("children"), list) else [])
        out.append(OutlineSection(
            title=x.get("title", ""),
            bullets=x.get("bullets", []),
            children=children,
            description=x.get("description"),
            word_count=x.get("word_count"),
            granularity=x.get("granularity")
        ))
    return out

def _section_to_dict(s: OutlineSection) -> dict:
    return {
        "title": s.title,
        "bullets": s.bullets or [],
        "children": [_section_to_dict(c) for c in (s.children or [])],
        "description": s.description,
        "word_count": s.word_count,
        "granularity": s.granularity
    }

def outline_for_text(text: str, type_name: str) -> List[OutlineSection]:
    client = get_client()
    settings = load_settings()
    model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    prompt = (
        "请为以下新闻稿提取此类新闻稿的通用结构化层级大纲，可参考用于后续同类型新闻稿生成(因此，大纲不得包含本内容的个性化内容，仅包含通用的结构化层级大纲和对应章节的内容描述介绍)，最多8个一级标题，每个标题可含3-6条要点。"
        "请特别注意提取每一部分的颗粒度（granularity）、大概字数范围（word_count）以及详细的写作指引（description）。"
        "仅返回一个JSON对象，字段sections为数组，元素包含："
        "title (标题), description (详细写作指引/描述), bullets (要点), word_count (字数范围, 如'100-150字'), granularity (颗粒度, 如'高/中/低'), 以及可选children（同结构）。"
        "章节命名参考：标题（主标题/副标题）、导语、事件背景、活动过程（开幕致辞与主题演讲/主旨分享与演讲）、"
        "重要讨论环节（圆桌讨论…）、合作签约仪式、专家与学者发言总结、论坛成果与未来展望、结语、媒体与活动支持、附加信息。"
        r"""
        下面是一个可参考的输出格式：
        {
            "sections": [
                {
                    "title": "标题",
                    "description": "生成简洁明了的标题，突出新闻事件的核心内容。",
                    "word_count": "20-50字",
                    "granularity": "高",
                    "bullets": [
                        "包含核心事件关键词",
                        "吸引读者注意"
                    ]
                },
                {
                    "title": "导语",
                    "description": "概括新闻稿的关键事件，包括谁、什么、何时、何地、为何。",
                    "word_count": "100-150字",
                    "granularity": "中",
                    "bullets": [
                        "时间、地点",
                        "主要人物/机构",
                        "核心事件概要"
                    ]
                },
                {
                    "title": "主体内容",
                    "description": "详细描述活动过程或事件细节。",
                    "word_count": "300-500字",
                    "granularity": "详细",
                    "bullets": [
                        "关键环节描述",
                        "重要数据或引用"
                    ]
                }
            ]
        }
        """
        "类型：" + type_name + "。内容：\n" + text[:6000]
    )
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        content = resp.choices[0].message.content or "{}"
        obj = _safe_json_object(content) or {"sections": []}
        arr = obj.get("sections") or []
        return _sections_from_json(arr)
    except Exception:
        lines = [l.strip() for l in text.splitlines() if l.strip()][:8]
        bullets = [l[:80] for l in lines[:6]]
        return [OutlineSection(title="内容概览", bullets=bullets)]

def merge_outlines_stream(outlines: List[List[OutlineSection]], type_name: str):
    client = get_client()
    settings = load_settings()
    model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    outlines_json = [[_section_to_dict(s) for s in o] for o in outlines]
    prompt = (
        "将多篇同类型新闻稿的大纲合并为一个统一可复用的通用层级大纲。"
        "若输入含具体人名、机构、日期、地点、数字、案例等，请通用化为描述性要点，不包含真实信息。"
        "仅返回一个JSON对象，字段sections为数组，元素含title、bullets、description、word_count、granularity以及可选children（同结构）。类型："
        + type_name + "。大纲集合：\n" + json.dumps(outlines_json, ensure_ascii=False)
    )
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        response_format={"type": "json_object"},
        stream=True
    )
    for chunk in resp:
        if hasattr(chunk, 'choices') and chunk.choices:
            content = chunk.choices[0].delta.content
            if content:
                yield content

def merge_outlines(outlines: List[List[OutlineSection]], type_name: str) -> List[OutlineSection]:
    # Fallback for non-streaming usage if any (currently extract uses this, we will replace it)
    full_text = ""
    for chunk in merge_outlines_stream(outlines, type_name):
        full_text += chunk
    
    try:
        obj = _safe_json_object(full_text) or {"sections": []}
        arr = obj.get("sections") or []
        return _sections_from_json(arr)
    except Exception:
        flat: List[OutlineSection] = []
        for o in outlines:
            flat.extend(o)
        return flat[:8]

def generate_article_stream(outline: Outline, materials: str, knowledge_base: Optional[Dict[str, str]] = None):
    client = get_client()
    settings = load_settings()
    model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    
    payload = {
        "type": outline.type_name,
        "outline": [_section_to_dict(s) for s in outline.sections],
        "materials": materials,
    }
    
    # Add knowledge base to payload if provided
    if knowledge_base:
        payload["knowledge_base"] = knowledge_base
    
    # Update prompt to include knowledge base instructions if available
    if knowledge_base:
        prompt = (
            "根据类型与统一大纲以及素材，撰写一篇完整的中文新闻稿用于发布在公众号平台，语气专业、结构严谨、可直接发布。输出不要包含任何额外的文字，只输出新闻稿内容。\n"
            "我提供了一个'knowledge_base'（人物知识库）。如果在撰写过程中提到了知识库中的人物（如演讲嘉宾、获奖者等），必须在第一次提及时，自然地融入其头衔、学历、职位和研究兴趣等信息。\n"
            "例如：'XXX教授（国家重点研发计划项目首席科学家、同济大学教授）在会上表示...'\n"
            "不要生硬地罗列简历，要结合语境自然插入。\n"
            "输出的新闻稿应该支持在建议放置图片或者视频素材的部分用特殊形式表示，比如【此处建议添加xxx图片素材】。\n"
            "遵循建议的层级大纲（包含children层级）和描述性要点、字数建议，保持专业和现代化表达，以实现宣传目的。输入：\n" + json.dumps(payload, ensure_ascii=False)
        )
    else:
        prompt = (
            "根据类型与统一大纲以及素材，撰写一篇完整的中文新闻稿用于发布在公众号平台，语气专业、结构严谨、可直接发布。输出不要包含任何额外的文字，只输出新闻稿内容。"
            "输出的新闻稿应该支持在建议放置图片或者视频素材的部分用特殊形式表示，比如【此处建议添加xxx图片素材】。\n"
            "遵循建议的层级大纲（包含children层级）和描述性要点、字数建议，保持专业和现代化表达，以实现宣传目的。输入：\n" + json.dumps(payload, ensure_ascii=False)
        )
    
    # 关键调试信息：给大模型的输入和最终prompt
    print("[DEBUG] 给大模型的输入（payload）：")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    print("[DEBUG] 给大模型的最终prompt：")
    print(prompt)
    print("[DEBUG] prompt长度：", len(prompt), "字符")

    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.6,
        stream=True
    )
    
    full_response = ""
    for chunk in resp:
        if hasattr(chunk, 'choices') and chunk.choices:
            content = chunk.choices[0].delta.content
            if content:
                full_response += content
                yield content
    

def generalize_sections(sections: List[OutlineSection], type_name: str) -> List[OutlineSection]:
    client = get_client()
    settings = load_settings()
    model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    src = [_section_to_dict(s) for s in sections]
    prompt = (
        "将以下层级大纲通用化，去除具体人名、机构、日期、地点、数字、案例等，仅保留描述性要点。"
        "仅返回一个JSON对象，字段sections为数组，元素含title、bullets、description、word_count、granularity以及可选children。类型："
        + type_name + "。输入：\n" + json.dumps({"sections": src}, ensure_ascii=False)
    )
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        content = resp.choices[0].message.content or "{}"
        obj = _safe_json_object(content) or {"sections": src}
        arr = obj.get("sections") or src
        return _sections_from_json(arr)
    except Exception:
        return sections

# Knowledge Base Management
KNOWLEDGE_FILES = {
    "people": DATA_DIR / "people.json",
    "brands": DATA_DIR / "brands.json",
    "courses": DATA_DIR / "courses.json",
    "research": DATA_DIR / "research.json",
    "alumni": DATA_DIR / "alumni.json",
}

def load_knowledge(category: str) -> Dict[str, str]:
    """Load knowledge base from file"""
    file_path = KNOWLEDGE_FILES.get(category)
    if not file_path or not file_path.exists():
        return {}
    try:
        data = json.loads(file_path.read_text(encoding="utf-8"))
        if category == "people" and "people" in data:
            return data["people"]
        return data
    except Exception:
        return {}

def load_people_meta() -> Dict[str, Dict[str, str]]:
    """Load people metadata (avatar, url)"""
    meta_path = DATA_DIR / "people_meta.json"
    if meta_path.exists():
        try:
            return json.loads(meta_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}

def save_people_meta(meta_data: Dict[str, Dict[str, str]]):
    """Save people metadata"""
    meta_path = DATA_DIR / "people_meta.json"
    meta_path.write_text(json.dumps(meta_data, ensure_ascii=False, indent=2), encoding="utf-8")

def save_knowledge(category: str, data: Dict[str, str]):
    """Save knowledge base to file and create a history backup"""
    file_path = KNOWLEDGE_FILES.get(category)
    if not file_path:
        raise ValueError(f"Unknown category: {category}")
    
    # Create history backup
    history_dir = DATA_DIR / "history" / category
    history_dir.mkdir(parents=True, exist_ok=True)
    timestamp = int(time.time())
    history_path = history_dir / f"{timestamp}.json"
    
    save_data = {"people": data} if category == "people" else data
    
    # Save backup
    if file_path.exists():
        shutil.copy2(file_path, history_path)
        
    # Save new data
    file_path.write_text(json.dumps(save_data, ensure_ascii=False, indent=2), encoding="utf-8")

@app.get("/api/knowledge/{category}/history")
def get_knowledge_history(category: str):
    """Get history versions for a category"""
    history_dir = DATA_DIR / "history" / category
    if not history_dir.exists():
        return []
    
    versions = []
    for file_path in history_dir.glob("*.json"):
        try:
            timestamp = int(file_path.stem)
            versions.append({
                "timestamp": timestamp,
                "time_str": time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(timestamp)),
                "file": file_path.name
            })
        except ValueError:
            pass
    return sorted(versions, key=lambda x: x["timestamp"], reverse=True)

@app.post("/api/knowledge/{category}/rollback/{timestamp}")
def rollback_knowledge(category: str, timestamp: int):
    """Rollback knowledge base to a specific version"""
    history_dir = DATA_DIR / "history" / category
    history_path = history_dir / f"{timestamp}.json"
    
    if not history_path.exists():
        raise HTTPException(status_code=404, detail="History version not found")
        
    file_path = KNOWLEDGE_FILES.get(category)
    if not file_path:
        raise HTTPException(status_code=400, detail=f"Unknown category: {category}")
        
    # Create a backup of current state before rollback
    current_backup_path = history_dir / f"{int(time.time())}_pre_rollback.json"
    if file_path.exists():
        shutil.copy2(file_path, current_backup_path)
        
    # Rollback
    shutil.copy2(history_path, file_path)
    return {"status": "success", "message": "Rollback successful"}

@app.get("/api/knowledge/{category}", response_model=Dict[str, str])
def get_knowledge(category: str):
    """Get all items from a knowledge category"""
    if category not in KNOWLEDGE_FILES:
        raise HTTPException(status_code=404, detail="Category not found")
    return load_knowledge(category)

@app.get("/api/people/meta", response_model=Dict[str, Dict[str, str]])
def get_people_meta_api():
    """Get metadata for people"""
    return load_people_meta()

@app.post("/api/knowledge/{category}")
def update_knowledge(category: str, data: Dict[str, str]):
    """Update all items in a knowledge category"""
    if category not in KNOWLEDGE_FILES:
        raise HTTPException(status_code=404, detail="Category not found")
    save_knowledge(category, data)
    return {"status": "success"}

import asyncio
import aiohttp

async def fetch_profile(session, name, title, profile_url, avatar_url, headers):
    try:
        async with session.get(profile_url, headers=headers, timeout=10) as profile_resp:
            profile_resp.raise_for_status()
            text = await profile_resp.text()
            profile_soup = BeautifulSoup(text, 'html.parser')
            
            intro_div = profile_soup.find('div', class_='saifProfessor_substance_r')
            if not intro_div:
                intro_div = profile_soup.find('div', class_='saifProfessor_substance')
                
            intro_text = ""
            if intro_div:
                p_tags = intro_div.find_all('p')
                if p_tags:
                    for p in p_tags:
                        p_text = p.get_text(strip=True)
                        if len(p_text) > 30:
                            intro_text += p_text + " "
                            if len(intro_text) > 300:
                                break
                                
                if not intro_text:
                    text_content = intro_div.get_text(separator=' ', strip=True)
                    if len(text_content) > 50:
                        intro_text = text_content[:500] + "..." if len(text_content) > 500 else text_content
            
            final_desc = intro_text.strip() if intro_text.strip() else title
            if final_desc != title:
                final_desc = f"{title}, {final_desc}"
                
            return name, final_desc, profile_url, avatar_url
    except Exception as e:
        print(f"Error scraping profile for {name}: {e}")
        return name, title, profile_url, avatar_url

@app.post("/api/knowledge/sync/faculty")
async def sync_faculty_knowledge():
    """Scrape faculty data from SAIF website and stream progress"""
    base_url = "https://www.saif.sjtu.edu.cn"
    categories = {
        "全职教授": "/faculty-research/full-time-professor",
        "访问教授": "/faculty-research/special-term-professor",
        "兼职教授": "/faculty-research/adjunct-professor",
        "兼聘教授": "/faculty-research/affiliated-professor"
    }
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    
    async def event_generator():
        try:
            # Load existing people to not overwrite manually added non-faculty people
            existing_people = load_knowledge("people")
            existing_meta = load_people_meta()
            new_people = {}
            new_meta = {}
            
            yield json.dumps({"status": "progress", "message": "正在获取教授列表..."}) + "\n"
            
            professors_to_fetch = []
            
            # Step 1: Get all links sequentially (this is fast)
            for category, path in categories.items():
                url = base_url + path
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.text, 'html.parser')
                items = soup.find_all('li', class_='saifProfessor_list2-item')
                
                for item in items:
                    link_tag = item.find('a')
                    if not link_tag or 'href' not in link_tag.attrs:
                        continue
                    
                    href = link_tag['href']
                    profile_url = href if href.startswith("http") else base_url + href
                    
                    img_tag = item.find('img')
                    avatar_url = ""
                    if img_tag and 'src' in img_tag.attrs:
                        src = img_tag['src']
                        avatar_url = src if src.startswith("http") else base_url + src

                    name_tag = item.find('p', class_='name')
                    title_tag = item.find('p', class_='desc')
                    
                    if name_tag and title_tag:
                        name = name_tag.get_text(strip=True)
                        title = title_tag.get_text(strip=True)
                        professors_to_fetch.append((name, title, profile_url, avatar_url))
            
            total = len(professors_to_fetch)
            yield json.dumps({"status": "progress", "message": f"共找到 {total} 位教授，开始并行获取详情..."}) + "\n"
            
            # Step 2: Fetch profiles in parallel
            completed = 0
            connector = aiohttp.TCPConnector(limit=10) # Limit concurrent connections
            async with aiohttp.ClientSession(connector=connector) as session:
                tasks = [fetch_profile(session, name, title, url, avatar_url, headers) for name, title, url, avatar_url in professors_to_fetch]
                
                for coro in asyncio.as_completed(tasks):
                    name, desc, prof_url, avatar_url = await coro
                    new_people[name] = desc
                    new_meta[name] = {
                        "url": prof_url,
                        "avatar": avatar_url
                    }
                    completed += 1
                    
                    # Update progress every 5 professors to avoid too many events
                    if completed % 5 == 0 or completed == total:
                        yield json.dumps({
                            "status": "progress", 
                            "message": f"正在处理进度 ({completed}/{total})...",
                            "progress": completed / total * 100
                        }) + "\n"
                                
            # Merge new people into existing (updating descriptions for scraped ones)
            existing_people.update(new_people)
            save_knowledge("people", existing_people)
            
            # Update meta
            existing_meta.update(new_meta)
            save_people_meta(existing_meta)
            
            yield json.dumps({
                "status": "success", 
                "message": f"成功同步了 {len(new_people)} 位教授的数据。",
                "count": len(new_people)
            }) + "\n"
            
        except Exception as e:
            yield json.dumps({"status": "error", "message": f"同步失败: {str(e)}"}) + "\n"

    return StreamingResponse(event_generator(), media_type="application/x-ndjson")

@app.get("/api/people", response_model=Dict[str, str])
def get_people():
    """Get all people from knowledge base (backward compatibility)"""
    return load_knowledge("people")

@app.post("/api/match-people", response_model=List[PersonMatchResult])
def match_people(req: PersonMatchRequest):
    """Match people names in materials using fuzzy matching with mixed similarity (pinyin + text)"""
    results: List[PersonMatchResult] = []
    materials = req.materials
    people_db = req.people_db
    
    # ========== 预处理：建立索引 ==========
    # 精确匹配集合
    db_name_set = set(people_db.keys())
    
    # 按姓氏（首字符）分组的倒排索引
    surname_index = {}
    for db_name in people_db:
        first_char = db_name[0] if db_name else ''
        if first_char not in surname_index:
            surname_index[first_char] = []
        surname_index[first_char].append(db_name)
    # ========== 预处理结束 ==========
    
    # 1. 提取候选词 (包含原素材里的错别字)
    # 例如：materials="程仕军开会", candidates=["程仕", "程仕军"]
    candidates = extract_candidates(materials)
    
    # Also extract English names
    words = re.findall(r'\b[A-Z][a-z]+\b', materials)
    i = 0
    while i < len(words):
        candidates.add(words[i])
        if i + 1 < len(words):
            combined = f"{words[i]} {words[i+1]}"
            candidates.add(combined)
        i += 1
    
    matched_db_names = set()  # 防止同一个人被多次添加
    all_matches = []  # 存储所有匹配结果，用于后续查找相似名字
    
    # 2. 【逻辑3：以库为主的变种】
    # 因为我们提取了候选词，现在拿每一个候选词去和数据库的所有人比对
    for cand in candidates:
        # ========== 精确匹配快速通道 ==========
        if cand in db_name_set and cand not in matched_db_names:
            # 精确命中，直接加入结果，跳过模糊匹配
            # 仍然查找相似名字供用户参考
            similar_names = []
            first_char = cand[0] if cand else ''
            group = surname_index.get(first_char, [])
            for other_name in group:
                if other_name != cand:
                    score = calculate_mixed_similarity(cand, other_name)
                    if score >= 0.5:
                        similar_names.append({'name': other_name, 'similarity': score})
            similar_names.sort(key=lambda x: x['similarity'], reverse=True)
            similar_names = similar_names[:5]

            all_matches.append({
                'match_type': 'exact',
                'material_name': cand,
                'db_name': cand,
                'similarity': 1.0,
                'db_info': people_db[cand],
                'similar_names': similar_names if similar_names else None
            })
            matched_db_names.add(cand)
            continue  # 跳过后续模糊匹配
        # ========== 精确匹配快速通道结束 ==========
        
        best_match_name = None
        best_score = 0.0
        all_similarities = []  # 存储所有相似度，用于查找相似名字
        
        # ========== 使用倒排索引缩小比对范围 ==========
        first_char = cand[0] if cand else ''
        is_chinese_cand = all('\u4e00' <= c <= '\u9fff' for c in cand)
        
        if is_chinese_cand:
            # 中文候选词：只和同姓的数据库名字比较
            candidates_to_compare = surname_index.get(first_char, [])
        else:
            # 英文名字：仍然遍历全库
            candidates_to_compare = list(people_db.keys())

        for db_name in candidates_to_compare:
            db_info = people_db[db_name]
            # 长度差异太大的不比对（例如 2个字 vs 4个字）
            if abs(len(cand) - len(db_name)) > 1:
                continue
            
            # 判断是否为中文名字
            is_chinese = is_chinese_cand and all('\u4e00' <= c <= '\u9fff' for c in db_name)
            
            if is_chinese:
                # 使用混合评分（拼音 + 字形）
                score = calculate_mixed_similarity(cand, db_name)
            else:
                # 英文名字使用纯文字相似度
                score = levenshtein_ratio(list(cand), list(db_name))
            
            all_similarities.append((db_name, score))
            
            if score > best_score:
                best_score = score
                best_match_name = db_name
        # ========== 倒排索引优化结束 ==========
        
        # 3. 阈值判定和相似名字查找
        # 降低阈值到 0.65，以便捕获更多可能的匹配（如"程仕军"vs"程四军"）
        # 同时降低相似名字的阈值到 0.5，以便显示更多可能的候选
        if best_match_name:
            # 查找所有相似名字（相似度 >= 0.5），用于显示可能的候选
            # 降低阈值以便显示更多可能的候选，帮助用户发现可能的拼写错误
            similar_names = []
            for other_db_name, other_score in all_similarities:
                if (other_db_name != best_match_name and 
                    other_score >= 0.5):
                    similar_names.append({
                        'name': other_db_name,
                        'similarity': other_score
                    })
            
            # 按相似度排序，保留前5个（增加数量以便用户看到更多选项）
            similar_names.sort(key=lambda x: x['similarity'], reverse=True)
            similar_names = similar_names[:5]
            
            # 如果最佳匹配达到阈值，直接添加
            if best_score >= 0.65:
                # 去重：如果这个数据库里的人还没被匹配过
                if best_match_name not in matched_db_names:
                    match_type = 'exact' if best_score >= 0.99 else 'fuzzy'
                    
                    all_matches.append({
                        'match_type': match_type,
                        'material_name': cand,
                        'db_name': best_match_name,
                        'similarity': round(best_score, 4),
                        'db_info': people_db[best_match_name],
                        'similar_names': similar_names if similar_names else None
                    })
                    
                    matched_db_names.add(best_match_name)
            # 如果最佳匹配低于阈值但 >= 0.55，且存在相似名字，也添加（标记为低置信度模糊匹配）
            # 这样即使相似度不是特别高，用户也能看到可能的候选
            elif best_score >= 0.55 and similar_names:
                if best_match_name not in matched_db_names:
                    all_matches.append({
                        'match_type': 'fuzzy',  # 标记为模糊匹配
                        'material_name': cand,
                        'db_name': best_match_name,
                        'similarity': round(best_score, 4),
                        'db_info': people_db[best_match_name],
                        'similar_names': similar_names if similar_names else None
                    })
                    
                    matched_db_names.add(best_match_name)
    
    # 按匹配度从高到低排序结果
    all_matches.sort(key=lambda x: x['similarity'], reverse=True)
    
    # 转换为 PersonMatchResult 格式
    for match in all_matches:
        # 转换 similar_names 为 SimilarName 对象列表
        similar_names_objects = None
        if match['similar_names']:
            similar_names_objects = [
                SimilarName(name=sn['name'], similarity=sn['similarity'])
                for sn in match['similar_names']
            ]
        
        results.append(PersonMatchResult(
            db_name=match['db_name'],
            material_name=match['material_name'],
            db_info=match['db_info'],
            match_type=match['match_type'],
            similarity=match['similarity'],
            similar_names=similar_names_objects
        ))
    
    return results

@app.post("/api/check-conflict", response_model=ConflictCheckResult)
def check_conflict(req: ConflictCheckRequest):
    """Check if material information conflicts with database information"""
    client = get_client()
    settings = load_settings()
    model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    
    prompt = (
        f"请阅读以下素材片段，提取人名 '{req.name}' 在文中的职位/身份信息。\n\n"
        f"素材片段：\n{req.material_context}\n\n"
        f"数据库中的信息：{req.db_info}\n\n"
        f"请执行以下任务：\n"
        f"1. 提取素材中关于 '{req.name}' 的职位/身份信息（即使信息不完整也要提取）\n"
        f"2. 判断素材中提取的职位/身份信息与数据库中的描述是否存在事实冲突？\n"
        f"3. 如果冲突，请给出具体的冲突证据（引用素材和数据库中的原文）\n"
        f"4. 给出你对判断结果的置信度（0到1之间的小数，1表示非常确定）\n\n"
        f"请以JSON格式返回，格式如下：\n"
        f'{{"has_conflict": true/false, '
        f'"material_info": "素材中提取的职位信息（必须返回，即使信息不完整）", '
        f'"reason": "冲突原因或一致性说明", '
        f'"evidence": {{"material_quote": "素材中的原文片段", "db_quote": "数据库中的对应信息"}}, '
        f'"confidence": 0.95}}'
    )
    
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        content = resp.choices[0].message.content or "{}"
        data = json.loads(content)
        
        # ========== 结果校验 ==========
        has_conflict = data.get("has_conflict", False)
        confidence = data.get("confidence", 0.5)
        evidence = data.get("evidence", {})
        
        # 校验1：如果声称有冲突但没有给出证据，降低置信度
        if has_conflict and not evidence.get("material_quote"):
            confidence = min(confidence, 0.3)
        
        # 校验2：如果声称有冲突但 reason 为空，降低置信度
        if has_conflict and not data.get("reason"):
            confidence = min(confidence, 0.3)
        
        # 校验3：置信度过低时，标记为不确定（而非直接报冲突）
        if has_conflict and confidence < 0.5:
            reason_suffix = "（置信度较低，建议人工核实）"
            data["reason"] = (data.get("reason", "") + reason_suffix).strip()
        # ========== 校验结束 ==========
        
        return ConflictCheckResult(
            has_conflict=has_conflict,
            material_info=data.get("material_info"),
            reason=data.get("reason"),
            confidence=round(confidence, 2)
        )
    except Exception as e:
        print(f"Error checking conflict: {e}")
        # 用 error 字段传递错误，而不是伪装成"无冲突"
        return ConflictCheckResult(
            has_conflict=False,
            reason=None,
            error=f"冲突检测失败: {str(e)}"
        )

@app.get("/api/types", response_model=List[TypeWithStats])
def get_types():
    types = load_types()
    ensure_dirs()
    result = []
    for t in types:
        count = 0
        for p in OUTLINES_DIR.glob(f"{t.id}*.json"):
            if p.name == f"{t.id}.json" or p.name.startswith(f"{t.id}__"):
                count += 1
        result.append(TypeWithStats(**t.model_dump(), template_count=count))
    return result

@app.post("/api/types", response_model=TypeItem)
def add_type(item: TypeItem):
    save_type(item)
    return item

@app.get("/api/outline/{type_id}", response_model=Outline)
def get_outline(type_id: str):
    path = outline_path(type_id, "default")
    if not path.exists():
        raise HTTPException(status_code=404, detail="outline not found")
    data = json.loads(path.read_text(encoding="utf-8"))
    if "template_id" not in data:
        data["template_id"] = "default"
    return Outline(**data)

@app.put("/api/outline/{type_id}", response_model=Outline)
def put_outline(type_id: str, outline: Outline):
    ensure_dirs()
    if outline.type_id != type_id:
        raise HTTPException(status_code=400, detail="type_id mismatch")
    path = outline_path(type_id, outline.template_id)
    path.write_text(json.dumps(outline.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
    return outline

@app.get("/api/outline/{type_id}/{template_id}", response_model=Outline)
def get_outline_with_template(type_id: str, template_id: str):
    path = outline_path(type_id, template_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="outline not found")
    data = json.loads(path.read_text(encoding="utf-8"))
    if "template_id" not in data:
        data["template_id"] = template_id
    return Outline(**data)

@app.put("/api/outline/{type_id}/{template_id}", response_model=Outline)
def put_outline_with_template(type_id: str, template_id: str, outline: Outline):
    ensure_dirs()
    if outline.type_id != type_id:
        raise HTTPException(status_code=400, detail="type_id mismatch")
    outline.template_id = template_id
    path = outline_path(type_id, template_id)
    path.write_text(json.dumps(outline.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
    return outline

@app.get("/api/outlines/{type_id}")
def list_outlines(type_id: str):
    ensure_dirs()
    results = []
    
    # 1. Check existing files
    for p in OUTLINES_DIR.glob(f"{type_id}*.json"):
        if p.name != f"{type_id}.json" and not p.name.startswith(f"{type_id}__"):
            continue
            
        try:
            # Use utf-8-sig to handle potential BOM
            data = json.loads(p.read_text(encoding="utf-8-sig"))
            meta = data.get("meta") or {}
            
            name = p.name
            if "__" in name:
                template_id = name.split("__", 1)[1][:-5]
            else:
                template_id = "default"
                
            results.append({
                "template_id": template_id,
                "meta": meta,
                "sections": data.get("sections", []),
                "sections_summary": [s["title"] for s in data.get("sections", [])[:5]]
            })
        except Exception as e:
            print(f"Error loading outline {p.name}: {e}")
            import traceback
            traceback.print_exc()
        
    return {"templates": results}

@app.post("/api/outline/{type_id}/{template_id}/favorite")
def toggle_favorite(type_id: str, template_id: str):
    path = outline_path(type_id, template_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="template not found")
    
    data = json.loads(path.read_text(encoding="utf-8"))
    meta = data.get("meta", {})
    meta["is_favorite"] = not meta.get("is_favorite", False)
    data["meta"] = meta
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "is_favorite": meta["is_favorite"]}

@app.delete("/api/outline/{type_id}/{template_id}")
def delete_template(type_id: str, template_id: str):
    path = outline_path(type_id, template_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="template not found")
    
    try:
        path.unlink()
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"failed to delete template: {str(e)}")

@app.post("/api/classify", response_model=TypeItem)
def classify(req: ClassifyRequest):
    types = load_types()
    return classify_type_for_texts(req.texts, types)

@app.post("/api/extract", response_model=Outline)
async def extract(files: List[UploadFile] = File(...), template_id: Optional[str] = Form("default")):
    ensure_dirs()
    if not files or len(files) < 1:
        raise HTTPException(status_code=400, detail="upload at least 1 file")

    # Eagerly read uploaded files into memory
    file_entries: List[tuple[str, bytes]] = []
    for f in files:
        name = f.filename.lower()
        b = await f.read()
        file_entries.append((name, b))

    # Parse texts
    texts: List[str] = []
    print("start parse files")
    for name, b in file_entries:
        if name.endswith(".txt"):
            texts.append(read_txt_bytes(b))
        elif name.endswith(".docx"):
            texts.append(read_docx_bytes(b))
        elif name.endswith(".pdf"):
            texts.append(read_pdf_bytes(b))

    # Classify type
    print("start classify type")
    types = load_types()
    selected = classify_type_for_texts(texts, types)

    # Extract outlines per file
    per_outlines: List[List[OutlineSection]] = []

    print("start extract")
    for t in texts:
        o = outline_for_text(t, selected.name)
        per_outlines.append(o)
        print(per_outlines)

    # Generalize sections
    # per_outlines = [generalize_sections(o, selected.name) for o in per_outlines]
    
    if len(per_outlines) == 1:
        merged_sections = per_outlines[0]
    else:
        existing_path = outline_path(selected.id, template_id)
        if existing_path.exists():
            try:
                existing = Outline(**json.loads(existing_path.read_text(encoding="utf-8")))
                per_outlines.append(generalize_sections(existing.sections, selected.name))
            except Exception:
                pass
        merged_sections = merge_outlines(per_outlines, selected.name)
    
    # Build and save outline
    outline = Outline(
        type_id=selected.id,
        type_name=selected.name,
        template_id=template_id or "default",
        sections=merged_sections
    )
    path = outline_path(selected.id, outline.template_id)
    try:
        path.write_text(json.dumps(outline.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        print(f"Error saving outline: {e}")

    return outline

@app.post("/api/files/parse")
async def parse_files(files: List[UploadFile] = File(...)):
    if not files or len(files) < 1:
        raise HTTPException(status_code=400, detail="upload at least 1 file")
    texts: List[str] = []
    for f in files:
        name = f.filename.lower()
        b = await f.read()
        await f.close()  # 显式关闭文件
        if name.endswith(".txt"):
            texts.append(read_txt_bytes(b))
        elif name.endswith(".docx"):
            texts.append(read_docx_bytes(b))
        elif name.endswith(".pdf"):
            texts.append(read_pdf_bytes(b))
        else:
            raise HTTPException(status_code=400, detail="unsupported file type")
    return {"texts": texts}

@app.put("/api/types/{type_id}", response_model=TypeItem)
def update_type(type_id: str, item: TypeItem):
    data = {"types": [x.model_dump() for x in load_types()]}
    found = False
    for t in data["types"]:
        if t["id"] == type_id:
            t["name"] = item.name
            found = True
            break
    if not found:
        raise HTTPException(status_code=404, detail="type not found")
    TYPES_REGISTRY_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return TypeItem(**item.model_dump())

@app.delete("/api/types/{type_id}")
def delete_type(type_id: str):
    types = [x.model_dump() for x in load_types()]
    new_types = [t for t in types if t["id"] != type_id]
    if len(new_types) == len(types):
        raise HTTPException(status_code=404, detail="type not found")
    TYPES_REGISTRY_FILE.write_text(json.dumps({"types": new_types}, ensure_ascii=False, indent=2), encoding="utf-8")
    for p in OUTLINES_DIR.glob(f"{type_id}*.json"):
        if p.name != f"{type_id}.json" and not p.name.startswith(f"{type_id}__"):
            continue
        try:
            p.unlink()
        except Exception:
            pass
    return {"ok": True}

@app.post("/api/generate")
def generate(req: GenerateRequest):
    import time
    import uuid
    article_id = str(uuid.uuid4())
    
    def generator():
        full_text = ""
        # Yield article ID first as a special event or just part of stream?
        # To simplify, we just stream text. The frontend can fetch history later if needed.
        # But we need to save it. So we accumulate.
        yield f"ID:{article_id}\n" 
        
        for chunk in generate_article_stream(req.outline, req.materials, req.knowledge_base):
            full_text += chunk
            yield chunk
            
        # Determine title
        if req.title and req.title.strip():
            final_title = req.title.strip()
        elif req.outline.meta and req.outline.meta.title:
            final_title = req.outline.meta.title
        else:
            final_title = f"{req.outline.type_name}稿件"

        # Save to history after generation is complete
        article = Article(
            id=article_id,
            title=final_title,
            type_id=req.outline.type_id,
            type_name=req.outline.type_name,
            template_id=req.outline.template_id or "default",
            content=full_text,
            materials=req.materials,
            created_at=time.time(),
            parent_id=None  # First generation has no parent
        )
        ensure_dirs()
        (ARTICLES_DIR / f"{article_id}.json").write_text(json.dumps(article.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")

    return StreamingResponse(generator(), media_type="text/plain")

@app.post("/api/modify-article")
def modify_article(req: ModifyRequest):
    import time
    import uuid
    
    # 1. 检查文章是否存在
    article_path = ARTICLES_DIR / f"{req.article_id}.json"
    if not article_path.exists():
        raise HTTPException(status_code=404, detail="Article not found")
    
    current_article = Article(**json.loads(article_path.read_text(encoding="utf-8")))
    
    # 2. 获取完整的历史链条 (V0 -> V1 -> V2 -> Current)
    history_chain = get_article_history(req.article_id)
    # 如果历史记录获取失败或为空，至少包含当前文章
    if not history_chain:
        history_chain = [current_article]
        
    def generator():
        client = get_client()
        settings = load_settings()
        model = settings.openai_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        
        # 3. 重新加载大纲 (Outline) 以构建最原始、最详细的 System Prompt
        # 我们需要知道这篇文章是用哪个 Type 和 Template 生成的
        try:
            # 尝试加载该文章对应的大纲
            outline_path_file = outline_path(current_article.type_id, current_article.template_id)
            if outline_path_file.exists():
                outline_data = json.loads(outline_path_file.read_text(encoding="utf-8"))
                sections = _sections_from_json(outline_data.get("sections", []))
            else:
                # 降级处理：使用通用大纲
                sections = _sections_from_json(GENERIC_OUTLINE_SECTIONS)
        except Exception:
            sections = _sections_from_json(GENERIC_OUTLINE_SECTIONS)

        # 4. 构建“详细版” System Prompt (包含 materials, outline, knowledge_base)
        # 注意：这里使用 req.knowledge_base 优先，如果没有则可以用 historic 的，但通常修改请求会带上 KB
        master_system_prompt = build_master_system_prompt(
            outline_sections=sections,
            type_name=current_article.type_name,
            materials=current_article.materials, # 使用文章原始素材
            knowledge_base=req.knowledge_base
        )
        
        # 5. 构建完整的对话历史 Messages
        messages = []
        messages.append({"role": "system", "content": master_system_prompt})
        
        # 遍历历史版本，重现对话流
        # 逻辑：V0 是 Assistant 的直接输出。V1 是 User(Query1) -> Assistant(V1) 的结果。
        for idx, art in enumerate(history_chain):
            if idx == 0:
                # 第一版文章，直接作为 Assistant 的回复（对应 System Prompt 的生成任务）
                messages.append({"role": "assistant", "content": art.content})
            else:
                # 后续版本，必须有 User 的修改指令
                # 如果 art.trigger_query 为空（旧数据），则根据内容伪造一个，否则使用真实的
                user_cmd = art.trigger_query if art.trigger_query else f"请基于上一版内容进行修改，参考内容：{history_chain[idx-1].content[:20]}..."
                
                messages.append({"role": "user", "content": user_cmd})
                messages.append({"role": "assistant", "content": art.content})
        
        # 6. 添加本次的修改请求
        messages.append({"role": "user", "content": req.user_query})
        
        # [DEBUG] 打印最终构建的 Context 供检查
        print("\n" + "="*50)
        print("[DEBUG] 深度修改模式 - 最终 Messages 结构：")
        for i, m in enumerate(messages):
            content_preview = m['content'][:100] + "..." if len(m['content']) > 100 else m['content']
            print(f"[{i}] {m['role']}: {content_preview}")
        print("="*50 + "\n")

        # 7. 调用大模型
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.6,
            stream=True
        )
        
        full_text = ""
        new_article_id = str(uuid.uuid4())
        yield f"ID:{new_article_id}\n"
        
        for chunk in resp:
            if hasattr(chunk, 'choices') and chunk.choices:
                content = chunk.choices[0].delta.content
                if content:
                    full_text += content
                    yield content
        
        # 8. 保存新版本 (关键：必须保存 trigger_query)
        new_title = current_article.title 
        if "（修改版）" not in new_title:
            new_title += "（修改版）"
            
        new_article = Article(
            id=new_article_id,
            title=new_title,
            type_id=current_article.type_id,
            type_name=current_article.type_name,
            template_id=current_article.template_id,
            content=full_text,
            materials=current_article.materials,
            created_at=time.time(),
            parent_id=req.article_id, # 链接父节点
            trigger_query=req.user_query # [新增] 保存本次的指令，供下次 Context 使用
        )
        
        new_article_path = ARTICLES_DIR / f"{new_article_id}.json"
        new_article_path.write_text(json.dumps(new_article.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
    
    return StreamingResponse(generator(), media_type="text/plain")

    
@app.get("/api/articles", response_model=List[Article])
def list_articles():
    ensure_dirs()
    articles = []
    for p in ARTICLES_DIR.glob("*.json"):
        try:
            articles.append(Article(**json.loads(p.read_text(encoding="utf-8"))))
        except Exception:
            pass
    return sorted(articles, key=lambda x: x.created_at, reverse=True)

@app.get("/api/articles/{article_id}", response_model=Article)
def get_article(article_id: str):
    path = ARTICLES_DIR / f"{article_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="article not found")
    return Article(**json.loads(path.read_text(encoding="utf-8")))

@app.get("/api/articles/{article_id}/history", response_model=List[Article])
def get_article_history_api(article_id: str):
    """API端点：获取文章的完整修改历史，包括原文章及其所有子文章"""
    return get_article_history(article_id)

@app.get("/api/articles/{article_id}/download/docx")
def download_docx(article_id: str):
    from io import BytesIO
    path = ARTICLES_DIR / f"{article_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="article not found")
    article = Article(**json.loads(path.read_text(encoding="utf-8")))
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    content = article.content
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
        else:
            clean_text = stripped.replace('**', '')
            p = doc.add_paragraph(clean_text)
    
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '黑体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    filename = f"{article.title}.docx"
    encoded_filename = quote(filename)
    
    return StreamingResponse(
        bio, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )

# @app.get("/api/articles/{article_id}/download/pdf")
# def download_pdf(article_id: str):
#     path = ARTICLES_DIR / f"{article_id}.json"
#     if not path.exists():
#         raise HTTPException(status_code=404, detail="article not found")
#     article = Article(**json.loads(path.read_text(encoding="utf-8")))
#     
#     font_registered = False
#     font_family = "Helvetica"
#     
#     font_candidates = [
#         "C:/Windows/Fonts/simhei.ttf",
#         "C:/Windows/Fonts/msyh.ttf",
#         "C:/Windows/Fonts/simkai.ttf",
#         "C:/Windows/Fonts/simfang.ttf",
#         "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
#         "/usr/share/fonts/truetype/noto/NotoSansSC-Regular.otf",
#     ]
#     
#     for font_path in font_candidates:
#         if os.path.exists(font_path):
#             try:
#                 pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
#                 font_family = 'ChineseFont'
#                 font_registered = True
#                 print(f"[PDF] Registered font from {font_path}")
#                 break
#             except Exception as e:
#                 print(f"[PDF] Failed to register {font_path}: {e}")
#                 continue
#     
#     if not font_registered:
#         print("[PDF] WARNING: No Chinese font found! PDF will show tofu blocks.")
#     
#     html_content = markdown.markdown(article.content)
#     
#     full_html = f"""
#     <html>
#     <head>
#     <meta charset="utf-8">
#     <style>
#         @page {{
#             size: A4;
#             margin: 2cm;
#         }}
#         body {{
#             font-family: "{font_family}";
#             font-size: 12px;
#             line-height: 1.8;
#             color: #333;
#         }}
#         h1 {{
#             text-align: center;
#             margin-bottom: 20px;
#             font-family: "{font_family}";
#             font-size: 18px;
#         }}
#         h2 {{
#             font-family: "{font_family}";
#             font-size: 15px;
#             margin-top: 15px;
#         }}
#         h3 {{
#             font-family: "{font_family}";
#             font-size: 13px;
#         }}
#         p {{ margin-bottom: 10px; }}
#     </style>
#     </head>
#     <body>
#         <h1>{article.title}</h1>
#         {html_content}
#     </body>
#     </html>
#     """
#     
#     result = BytesIO()
#     pdf = pisa.CreatePDF(BytesIO(full_html.encode("utf-8")), result, encoding='utf-8')
#     
#     if pdf.err:
#         raise HTTPException(status_code=500, detail="PDF generation failed")
#         
#     result.seek(0)
#     filename = f"{article.title}.pdf"
#     encoded_filename = quote(filename)
#     
#     return StreamingResponse(
#         result, 
#         media_type="application/pdf",
#         headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
#     )



@app.put("/api/articles/{article_id}", response_model=Article)
def update_article(article_id: str, article: Article):
    path = ARTICLES_DIR / f"{article_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="article not found")
    if article.id != article_id:
        raise HTTPException(status_code=400, detail="id mismatch")
    
    path.write_text(json.dumps(article.model_dump(), ensure_ascii=False, indent=2), encoding="utf-8")
    return article

@app.delete("/api/articles/{article_id}")
def delete_article(article_id: str):
    path = ARTICLES_DIR / f"{article_id}.json"
    if path.exists():
        path.unlink()
    return {"ok": True}

@app.get("/api/settings", response_model=Settings)
def get_settings():
    return load_settings()

@app.post("/api/settings", response_model=Settings)
def update_settings(s: Settings):
    save_settings(s)
    return s

if __name__ == "__main__":
    ensure_dirs()
    import uvicorn
    host = os.getenv("SERVER_HOST", "127.0.0.1")
    port = int(os.getenv("SERVER_PORT", "8002"))
    uvicorn.run(app, host=host, port=port)
